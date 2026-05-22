#!/usr/bin/env python3
"""Build TRC-Audit.docx — the audit team's Word document.

Walks the Obsidian vault, stitches the master + compliance posture + every
register and per-asset note into a single document for the audit team.
"""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from vault import (
    VAULT_ROOT,
    ASSET_FOLDERS,
    COMPLIANCE_FOLDERS,
    load_assets,
    load_compliance,
    load_root_note,
)

OUT_DOCX = VAULT_ROOT / "TRC-Audit.docx"

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
LIST_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
CHECKBOX_RE = re.compile(r"^\s*-\s*\[( |x|X)\]\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
WIKILINK_RE = re.compile(r"\[\[([^\]|]+)(?:\|([^\]]+))?\]\]")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s\-:|]+\|?\s*$")


def strip_wikilinks(text: str) -> str:
    def repl(m):
        return m.group(2) or m.group(1)
    return WIKILINK_RE.sub(repl, text)


def add_runs(paragraph, text: str) -> None:
    text = strip_wikilinks(text)
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _add_italic_runs(paragraph, text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        _add_italic_runs(paragraph, text[pos:])


def _add_italic_runs(paragraph, text: str) -> None:
    pos = 0
    for m in ITALIC_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_markdown(doc: Document, md_text: str) -> None:
    in_code = False
    table_buffer: list[list[str]] = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        if len(table_buffer) < 2:
            for row in table_buffer:
                p = doc.add_paragraph(" | ".join(row))
            table_buffer = []
            return
        header = table_buffer[0]
        rows = table_buffer[1:]
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        table.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, h.strip())
            for run in p.runs:
                run.bold = True
        for i, row in enumerate(rows):
            for j in range(len(header)):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                add_runs(cell.paragraphs[0], (row[j] if j < len(row) else "").strip())
        table_buffer = []

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if line.strip().startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if line.lstrip().startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if TABLE_SEP_RE.match(line):
                continue
            table_buffer.append(cells)
            continue
        else:
            flush_table()

        if not line.strip():
            doc.add_paragraph()
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            doc.add_heading(strip_wikilinks(m.group(2).strip()), level=min(level, 4))
            continue

        m = CHECKBOX_RE.match(line)
        if m:
            mark = "[x]" if m.group(1).lower() == "x" else "[ ]"
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{mark} ")
            run.bold = True
            add_runs(p, m.group(2))
            continue

        m = LIST_RE.match(line)
        if m:
            indent = len(m.group(1)) // 2
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, m.group(2))
            continue

        m = QUOTE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(strip_wikilinks(m.group(1)))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        p = doc.add_paragraph()
        add_runs(p, line)

    flush_table()


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TRC Master Audit File")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Maintained by Judge Dredd on behalf of the Governor")
    run.italic = True
    run.font.size = Pt(12)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Jurisdiction: EU (Malta) — GDPR + EU AI Act")
    run.font.size = Pt(11)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(f"Generated: {date.today().isoformat()}")

    doc.add_paragraph()
    doc.add_paragraph(
        "This document is the single source of truth for what TRC is running, who "
        "owns each piece, which third-party tools touch TRC personal data, and "
        "how TRC meets each applicable EU obligation. It is regenerated in full "
        "every time an agent, job, project, dataset, integration, or compliance "
        "artefact is added, changed, or retired. The Alfred / Polymarket bot is "
        "intentionally excluded from this audit."
    )
    doc.add_page_break()


def section_break(doc: Document, title: str) -> None:
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def main() -> int:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc)

    # 1. Compliance Posture first — the audit team reads this front.
    posture = load_root_note("Compliance Posture.md")
    if posture:
        render_markdown(doc, posture)

    # 2. Master Audit
    section_break(doc, "Master Audit")
    master = load_root_note("Master Audit.md")
    if master:
        render_markdown(doc, master)

    # 3. Evidence Plan — per project, per control, exact artefacts required
    plan_path = VAULT_ROOT / "Compliance" / "Evidence Plan.md"
    if plan_path.exists():
        section_break(doc, "Evidence Plan")
        render_markdown(doc, plan_path.read_text(encoding="utf-8"))

    # 3. Compliance artefacts, folder by folder
    compliance = load_compliance()
    for folder in COMPLIANCE_FOLDERS:
        notes = compliance.get(folder, [])
        if not notes:
            continue
        section_break(doc, folder.split("/", 1)[-1])
        for n in notes:
            doc.add_heading(n.title, level=2)
            render_markdown(doc, n.body)

    # 4. Asset registers
    assets = load_assets()
    for folder in ASSET_FOLDERS:
        notes = assets.get(folder, [])
        if not notes:
            continue
        section_break(doc, folder)
        for n in notes:
            doc.add_heading(n.title, level=2)
            render_markdown(doc, n.body)

    # 5. Open Questions + Change Log at the end
    section_break(doc, "Open Questions")
    render_markdown(doc, load_root_note("Open Questions.md"))

    section_break(doc, "Change Log")
    render_markdown(doc, load_root_note("Change Log.md"))

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
